import pandas as pd
import statsmodels.api as sm
from statsmodels.formula.api import ols
from statsmodels.multivariate.manova import MANOVA
from docx import Document

# Load the data
file_path = "anonymized_data.xlsx"
df = pd.read_excel(file_path)
df = df.rename(columns={'Type of test': 'IV'})

# Formatter
def fmt(val, places=2):
    return f"{val:.{places}f}"

# Run MANOVA
manova = MANOVA.from_formula('Anxiety + Stress + Spirituality ~ IV', data=df)
summary = str(manova.mv_test())
lines = summary.splitlines()
iv_start = next(i for i, line in enumerate(lines) if 'IV' in line and 'Value' in line)
iv_section = lines[iv_start + 1 : iv_start + 5]

# Build Word Document
doc = Document()
doc.add_heading('MANCOVA', level=1)

# Multivariate Table
doc.add_heading('Multivariate Tests', level=2)
multivar_table = doc.add_table(rows=1, cols=6)
hdr_cells = multivar_table.rows[0].cells
hdr_cells[0].text = 'Type of Test'
hdr_cells[1].text = 'value'
hdr_cells[2].text = 'F'
hdr_cells[3].text = 'df1'
hdr_cells[4].text = 'df2'
hdr_cells[5].text = 'p'

for line in iv_section:
    parts = [p for p in line.strip().split() if p]
    if len(parts) >= 6:
        test_name = " ".join(parts[:-5])
        value, fval, df1, df2, pval = parts[-5:]
        row_cells = multivar_table.add_row().cells
        row_cells[0].text = test_name
        row_cells[1].text = fmt(float(value), 4)
        row_cells[2].text = fmt(float(fval), 2)
        row_cells[3].text = str(int(float(df1)))
        row_cells[4].text = str(int(float(df2)))
        row_cells[5].text = fmt(float(pval), 3)

# Univariate Tests
doc.add_heading('Univariate Tests', level=2)
uni_table = doc.add_table(rows=1, cols=6)
uni_table.style = 'Table Grid'
hdr = uni_table.rows[0].cells
hdr[0].text = 'Dependent Variable'
hdr[1].text = 'Sum of Squares'
hdr[2].text = 'df'
hdr[3].text = 'Mean Square'
hdr[4].text = 'F'
hdr[5].text = 'p'

for dv in ['Anxiety', 'Spirituality', 'Stress']:
    model = ols(f'{dv} ~ C(IV)', data=df).fit()
    anova = sm.stats.anova_lm(model, typ=2)
    ss_effect = anova.loc['C(IV)', 'sum_sq']
    df_effect = anova.loc['C(IV)', 'df']
    ms_effect = ss_effect / df_effect
    f_val = anova.loc['C(IV)', 'F']
    p_val = anova.loc['C(IV)', 'PR(>F)']
    row = uni_table.add_row().cells
    row[0].text = dv
    row[1].text = fmt(ss_effect, 0)
    row[2].text = str(int(df_effect))
    row[3].text = fmt(ms_effect, 1)
    row[4].text = fmt(f_val, 2)
    row[5].text = fmt(p_val, 3)

# Residuals Table
doc.add_heading('Residuals', level=2)
resid_table = doc.add_table(rows=1, cols=4)
resid_table.style = 'Table Grid'
hdr = resid_table.rows[0].cells
hdr[0].text = 'Dependent Variable'
hdr[1].text = 'Sum of Squares'
hdr[2].text = 'df'
hdr[3].text = 'Mean Square'

for dv in ['Anxiety', 'Spirituality', 'Stress']:
    model = ols(f'{dv} ~ C(IV)', data=df).fit()
    anova = sm.stats.anova_lm(model, typ=2)
    ss_resid = anova.loc['Residual', 'sum_sq']
    df_resid = anova.loc['Residual', 'df']
    ms_resid = ss_resid / df_resid
    row = resid_table.add_row().cells
    row[0].text = dv
    row[1].text = fmt(ss_resid, 0)
    row[2].text = str(int(df_resid))
    row[3].text = fmt(ms_resid, 1)

# Save the document
output_path = "mancova_output.docx"
doc.save(output_path)
output_path
